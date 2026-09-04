import os
import base64
import csv
import io
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import Dict, Any, Optional

load_dotenv()

from src.orchestrator import Orchestrator

app = FastAPI(
    title="ModulAI - AI Core FastAPI",
    description="Moteur générique d'exécution IA pour ModulAI.",
    version="1.0.0"
)

orchestrator = Orchestrator()

from typing import Dict, Any, Optional, List

class ExecutePayload(BaseModel):
    module: str
    use_case: str
    system_prompt_template: Optional[str] = None
    user_prompt: Optional[str] = None
    variables: Optional[Dict[str, Any]] = None
    output_schema: Optional[List[Dict[str, Any]]] = None
    rag_config: Optional[Dict[str, Any]] = None
    model_options: Optional[Dict[str, Any]] = None
    project_id: Optional[str] = None
    project_name: Optional[str] = None
    module_id: Optional[str] = None
    knowledge_base_id: Optional[str] = None
    knowledge_base_ids: Optional[List[str]] = None
    configuration: Optional[Dict[str, Any]] = None
    input_reference: Optional[Dict[str, Any]] = None
    context_reference: Optional[Dict[str, Any]] = None

class RagIndexPayload(BaseModel):
    collection: str
    documents: List[str]
    ids: List[str]
    metadatas: Optional[List[Dict[str, Any]]] = None

class RagSearchPayload(BaseModel):
    collection: str
    query: str
    top_k: int = 5
    filter_metadata: Optional[Dict[str, Any]] = None

class RagInspectPayload(BaseModel):
    collection: str
    limit: int = 100
    offset: int = 0

@app.get("/")
def read_root():
    return {"status": "online", "service": "ModulAI Core", "version": "1.0.0"}

@app.post("/api/rag/extract")
def extract_rag_file(payload: Dict[str, Any]):
    filename = str(payload.get("filename", ""))
    encoded = payload.get("content_base64")
    if not filename or not encoded:
        raise HTTPException(status_code=400, detail="filename and content_base64 are required")
    try:
        raw = base64.b64decode(encoded)
        extension = os.path.splitext(filename.lower())[1]
        if extension in {".txt", ".md", ".text", ".csv"}:
            text = raw.decode("utf-8-sig", errors="replace")
            if extension == ".csv":
                rows = list(csv.reader(io.StringIO(text)))
                text = "\n".join(" | ".join(row) for row in rows)
            extractor = "text-csv" if extension == ".csv" else "text"
        elif extension == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
            extractor = "pypdf"
        elif extension == ".docx":
            from docx import Document
            document = Document(io.BytesIO(raw))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            tables = [" | ".join(cell.text.strip() for cell in row.cells) for table in document.tables for row in table.rows]
            text = "\n".join(paragraphs + tables)
            extractor = "python-docx"
        elif extension == ".xlsx":
            from openpyxl import load_workbook
            workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            sections = []
            for sheet in workbook.worksheets:
                sections.append(f"Feuille: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(value) for value in row if value is not None]
                    if values:
                        sections.append(" | ".join(values))
            text = "\n".join(sections)
            extractor = "openpyxl"
        else:
            raise HTTPException(status_code=415, detail="Supported file types: PDF, DOCX, XLSX, CSV, TXT, MD")
        if not text.strip():
            raise HTTPException(status_code=422, detail="The extracted document is empty")
        return {"filename": filename, "type": extension.lstrip("."), "extractor": extractor, "text": text[:1000000]}
    except HTTPException:
        raise
    except ImportError as error:
        raise HTTPException(status_code=503, detail=f"Extractor dependency missing: {error.name}")
    except Exception as error:
        raise HTTPException(status_code=422, detail=f"File extraction failed: {error}")

@app.post("/api/rag/index")
def index_rag_documents(payload: RagIndexPayload):
    if not payload.documents:
        raise HTTPException(status_code=400, detail="documents must not be empty")
    if len(payload.documents) != len(payload.ids):
        raise HTTPException(status_code=400, detail="documents and ids must have the same length")
    metadatas = payload.metadatas or [{} for _ in payload.documents]
    if len(metadatas) != len(payload.documents):
        raise HTTPException(status_code=400, detail="metadatas and documents must have the same length")
    try:
        store = orchestrator.get_vector_store(payload.collection)
        store.upsert_documents(payload.documents, metadatas, payload.ids)
        return {"status": "indexed", "collection": payload.collection, "count": len(payload.documents), "total": store.count()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"RAG indexing failed: {e}")

@app.post("/api/rag/search")
def search_rag_documents(payload: RagSearchPayload):
    try:
        store = orchestrator.get_vector_store(payload.collection)
        results = store.search(payload.query, top_k=max(1, min(payload.top_k, 20)), filter_metadata=payload.filter_metadata)
        return {"collection": payload.collection, "query": payload.query, "results": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"RAG search failed: {e}")

@app.post("/api/rag/inspect")
def inspect_rag_collection(payload: RagInspectPayload):
    try:
        store = orchestrator.get_vector_store(payload.collection)
        return store.inspect(payload.limit, payload.offset)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"RAG inspection failed: {e}")

@app.post("/api/config/resolve")
def resolve_configuration(payload: ExecutePayload):
    """Prévisualise la configuration effective sans appeler de modèle ni modifier l’audit."""
    try:
        return orchestrator.config_resolver.resolve(payload.dict())["snapshot"]
    except Exception as error:
        raise HTTPException(status_code=500, detail=str(error))

@app.post("/api/execute")
def execute_use_case(payload: ExecutePayload):
    try:
        result = orchestrator.execute(payload.dict())
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host=os.getenv("AI_CORE_HOST", "0.0.0.0"),
        port=int(os.getenv("AI_CORE_PORT", "8001")),
        reload=os.getenv("AI_CORE_RELOAD", "true").lower() == "true",
    )
